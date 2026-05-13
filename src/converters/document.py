# Histórico de Alterações:
#   Data       | Desenvolvedor  | Ticket | Observação
#   -----------|----------------|--------|-------------------
#   2026-05-13 | AFSIS Dev      | #001   | Implementação v1.1.0


import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Callable, Optional

from src.converters.base import BaseConverter, ConversionError
from src.utils.logger import get_logger

logger = get_logger(__name__)


def _find_libreoffice() -> Optional[str]:
    for name in ["libreoffice", "soffice"]:
        path = shutil.which(name)
        if path:
            return path
    # Check common install locations on Linux
    for loc in [
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/opt/libreoffice/program/soffice",
        "/snap/bin/libreoffice",
    ]:
        if Path(loc).exists():
            return loc
    return None


def _libreoffice_convert(input_path: Path, output_dir: Path, target_format: str) -> Path:
    libreoffice = _find_libreoffice()
    if not libreoffice:
        raise ConversionError(
            "LibreOffice não encontrado. Para converter documentos, instale o LibreOffice:\n"
            "  Ubuntu/Debian: sudo apt install libreoffice\n"
            "  Fedora: sudo dnf install libreoffice\n"
            "  macOS: brew install libreoffice\n"
            "  Windows: https://www.libreoffice.org/download/"
        )

    cmd = [
        libreoffice,
        "--headless",
        "--convert-to",
        target_format,
        "--outdir",
        str(output_dir),
        str(input_path),
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            raise ConversionError(f"LibreOffice falhou: {result.stderr.strip()}")

        # Find the output file
        expected_name = input_path.stem + "." + target_format
        output_path = output_dir / expected_name
        if output_path.exists():
            return output_path

        # Try matching files
        for f in output_dir.iterdir():
            if f.stem == input_path.stem and f.suffix.lower() == f".{target_format}":
                return f

        raise ConversionError("Arquivo de saída do LibreOffice não encontrado")
    except subprocess.TimeoutExpired:
        raise ConversionError("Conversão via LibreOffice excedeu o tempo limite (120s)")
    except FileNotFoundError:
        raise ConversionError("LibreOffice não encontrado no sistema")


class DocumentConverter(BaseConverter):

    def supported_input_formats(self) -> list[str]:
        return [".docx", ".odt", ".txt", ".rtf", ".html", ".htm"]

    def supported_output_formats(self) -> list[str]:
        return [".pdf", ".docx", ".txt", ".html"]

    def category(self) -> str:
        return "Documento"

    def get_options_schema(self) -> dict:
        return {
            "encoding": {"type": "str", "default": "utf-8", "label": "Encoding texto"},
        }

    def convert(
        self,
        input_path: Path,
        output_path: Path,
        progress_callback: Optional[Callable[[float], None]] = None,
    ) -> None:
        ext_in = input_path.suffix.lower()
        ext_out = output_path.suffix.lower()

        if not self.can_convert(ext_in, ext_out):
            raise ConversionError(f"Não é possível converter {ext_in} → {ext_out}")

        # DOCX → DOCX is a no-op copy
        if ext_in == ".docx" and ext_out == ".docx":
            shutil.copy2(input_path, output_path)
            if progress_callback:
                progress_callback(1.0)
            return

        # TXT → TXT is a copy
        if ext_in == ".txt" and ext_out == ".txt":
            shutil.copy2(input_path, output_path)
            if progress_callback:
                progress_callback(1.0)
            return

        # DOCX → TXT: extract text
        if ext_in == ".docx" and ext_out == ".txt":
            self._docx_to_txt(input_path, output_path, progress_callback)
            return

        # TXT → DOCX: create document from text
        if ext_in == ".txt" and ext_out == ".docx":
            self._txt_to_docx(input_path, output_path, progress_callback)
            return

        # RTF → TXT: extract text
        if ext_in == ".rtf" and ext_out == ".txt":
            self._rtf_to_txt(input_path, output_path, progress_callback)
            return

        # HTML → TXT: strip tags
        if ext_in in (".html", ".htm") and ext_out == ".txt":
            self._html_to_txt(input_path, output_path, progress_callback)
            return

        # For conversions requiring LibreOffice (ANY → PDF, ODT → *, RTF → PDF/DOCX, HTML → DOCX/PDF)
        self._convert_via_libreoffice(input_path, output_path, ext_out, progress_callback)

    def _docx_to_txt(
        self,
        input_path: Path,
        output_path: Path,
        progress_callback: Optional[Callable[[float], None]] = None,
    ) -> None:
        try:
            from docx import Document
        except ImportError:
            raise ConversionError("python-docx não está instalado. Instale com: pip install python-docx")

        try:
            if progress_callback:
                progress_callback(0.3)
            doc = Document(str(input_path))
            paragraphs = [p.text for p in doc.paragraphs]
            if progress_callback:
                progress_callback(0.7)
            encoding = self.options.get("encoding", "utf-8")
            with open(output_path, "w", encoding=encoding) as f:
                f.write("\n".join(paragraphs))
            if progress_callback:
                progress_callback(1.0)
        except Exception as e:
            raise ConversionError(f"Erro ao extrair texto do DOCX: {e}")

    def _txt_to_docx(
        self,
        input_path: Path,
        output_path: Path,
        progress_callback: Optional[Callable[[float], None]] = None,
    ) -> None:
        try:
            from docx import Document
        except ImportError:
            raise ConversionError("python-docx não está instalado. Instale com: pip install python-docx")

        try:
            if progress_callback:
                progress_callback(0.3)
            encoding = self.options.get("encoding", "utf-8")
            with open(input_path, "r", encoding=encoding) as f:
                lines = f.read().splitlines()
            doc = Document()
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line)
            if progress_callback:
                progress_callback(0.7)
            doc.save(str(output_path))
            if progress_callback:
                progress_callback(1.0)
        except Exception as e:
            raise ConversionError(f"Erro ao criar DOCX: {e}")

    def _rtf_to_txt(
        self,
        input_path: Path,
        output_path: Path,
        progress_callback: Optional[Callable[[float], None]] = None,
    ) -> None:
        # Use LibreOffice as fallback to convert RTF → TXT
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_p = Path(tmpdir)
            result = _libreoffice_convert(input_path, tmpdir_p, "txt")
            shutil.move(str(result), str(output_path))
        if progress_callback:
            progress_callback(1.0)

    def _html_to_txt(
        self,
        input_path: Path,
        output_path: Path,
        progress_callback: Optional[Callable[[float], None]] = None,
    ) -> None:
        try:
            if progress_callback:
                progress_callback(0.3)
            encoding = self.options.get("encoding", "utf-8")
            with open(input_path, "r", encoding=encoding) as f:
                html = f.read()

            try:
                from html.parser import HTMLParser

                class TextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text: list[str] = []

                    def handle_data(self, data):
                        t = data.strip()
                        if t:
                            self.text.append(t)

                extractor = TextExtractor()
                extractor.feed(html)

                with open(output_path, "w", encoding=encoding) as f:
                    f.write("\n".join(extractor.text))
            except Exception:
                # Simple strip tags fallback
                import re
                text = re.sub(r"<[^>]+>", "\n", html)
                text = re.sub(r"\n+", "\n", text).strip()
                with open(output_path, "w", encoding=encoding) as f:
                    f.write(text)

            if progress_callback:
                progress_callback(1.0)
        except Exception as e:
            raise ConversionError(f"Erro ao extrair texto do HTML: {e}")

    def _convert_via_libreoffice(
        self,
        input_path: Path,
        output_path: Path,
        target_ext: str,
        progress_callback: Optional[Callable[[float], None]] = None,
    ) -> None:
        if progress_callback:
            progress_callback(0.2)

        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_p = Path(tmpdir)
            result = _libreoffice_convert(
                input_path, tmpdir_p, target_ext.lstrip(".")
            )
            if progress_callback:
                progress_callback(0.8)

            if result.suffix.lower() != output_path.suffix.lower():
                # Rename to match expected output
                final = result.rename(output_path)
            else:
                shutil.move(str(result), str(output_path))

            if progress_callback:
                progress_callback(1.0)
