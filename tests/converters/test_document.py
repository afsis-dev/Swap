# Histórico de Alterações:
#   Data       | Desenvolvedor  | Ticket | Observação
#   -----------|----------------|--------|-------------------
#   2026-05-13 | AFSIS Dev      | #001   | Implementação v1.1.0


import pytest
import tempfile
from pathlib import Path
from src.converters.document import DocumentConverter


@pytest.fixture
def tmp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_docx(tmp_dir):
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Second paragraph")
        path = tmp_dir / "test.docx"
        doc.save(str(path))
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


class TestDocumentConverter:

    def test_docx_to_txt(self, sample_docx, tmp_dir):
        converter = DocumentConverter()
        output = tmp_dir / "output.txt"
        converter.convert(sample_docx, output)
        assert output.exists()
        content = output.read_text()
        assert "Hello World" in content

    def test_txt_to_docx(self, tmp_dir):
        txt_path = tmp_dir / "input.txt"
        txt_path.write_text("Line 1\nLine 2\nLine 3", encoding="utf-8")

        converter = DocumentConverter()
        output = tmp_dir / "output.docx"
        converter.convert(txt_path, output)
        assert output.exists()

    def test_supported_formats(self):
        converter = DocumentConverter()
        assert ".docx" in converter.supported_input_formats()
        assert ".odt" in converter.supported_input_formats()
        assert ".txt" in converter.supported_input_formats()
        assert ".pdf" in converter.supported_output_formats()
        assert converter.category() == "Documento"

    def test_docx_to_docx_copy(self, sample_docx, tmp_dir):
        converter = DocumentConverter()
        output = tmp_dir / "output.docx"
        converter.convert(sample_docx, output)
        assert output.exists()
